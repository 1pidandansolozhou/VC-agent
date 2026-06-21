from typing import List
from docx import Document

from exporters.summary import weekly_summary
from models.schema import Deal


def write_word(deals: List[Deal], date_range: str, docx_path: str) -> str:
    """周报 Word：剔除 stale（旧闻），严格按第4节结构渲染。"""
    keep = [d for d in deals if d.date_status != "stale"]

    doc = Document()
    doc.add_heading(f"AI 创业资讯 {date_range}", 0)

    doc.add_heading("本周综述", 1)
    doc.add_paragraph(weekly_summary(keep))

    doc.add_heading("一、早期项目", 1)
    rank = {"high": 0, "mid": 1, "low": 2}
    for d in sorted(keep, key=lambda x: rank.get(x.importance, 1)):
        p = doc.add_paragraph()
        r = p.add_run(d.title)
        r.bold = True

        if d.official_site:
            doc.add_paragraph(f"官方网站：{d.official_site}")

        parts = [d.round]
        if d.amount and d.amount != "未披露":
            parts.append(f"融资 {d.amount}")
        if d.valuation and d.valuation != "未披露":
            parts.append(f"估值 {d.valuation}")
        if d.investors:
            parts.append(d.investors)
        if d.region:
            parts.append(d.region)
        if d.verified_date:
            parts.append(f"公布 {d.verified_date}")
        doc.add_paragraph(" · ".join(parts))

        doc.add_paragraph(d.detail)

        if d.team:
            doc.add_paragraph(f"团队：{d.team}")
        if d.business:
            doc.add_paragraph(f"业务：{d.business}")

        doc.add_paragraph("")

    doc.save(docx_path)
    return docx_path
